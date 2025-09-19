import os
import google.generativeai as genai
import docx
from pypdf import PdfReader
import re

# -----------------------------
# Funciones auxiliares
# -----------------------------
def extract_text_from_file(file_path):
    """Extrae texto de archivos .docx o .pdf."""
    if file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_path.endswith('.pdf'):
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
    else:
        raise ValueError("Formato de archivo no soportado. Usa .docx o .pdf.")

def split_document_into_chunks(text, max_chunk_size=3000):
    """Divide el documento en chunks manejables."""
    # Primero intentar dividir por secciones/capÃ­tulos
    sections = re.split(r'\n\s*(?:[0-9]+\.|\b(?:CAPÃTULO|SECCIÃ“N|MÃ“DULO|FUNCIONALIDAD)\b)', text, flags=re.IGNORECASE)

    chunks = []
    current_chunk = ""

    for section in sections:
        if len(current_chunk) + len(section) < max_chunk_size:
            current_chunk += section
        else:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            current_chunk = section

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    # Si no hay divisiones claras, dividir por pÃ¡rrafos
    if len(chunks) == 1 and len(text) > max_chunk_size:
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""

        for para in paragraphs:
            if len(current_chunk) + len(para) < max_chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

    return chunks

def create_analysis_prompt(document_text, role, business_context=None):
    """Crea un prompt inicial para anÃ¡lisis de funcionalidades."""
    context_section = ""
    if business_context and business_context.strip():
        context_section = f"""
CONTEXTO ADICIONAL DE NEGOCIO:
{business_context}

IMPORTANTE: Debes tomar en cuenta TANTO los requerimientos del documento COMO el contexto adicional proporcionado.
"""

    return f"""
Eres un analista de negocios Senior. Tu tarea es IDENTIFICAR Y LISTAR todas las funcionalidades del siguiente documento.

DOCUMENTO A ANALIZAR:
{document_text}
{context_section}
INSTRUCCIONES:
1. Lee COMPLETAMENTE el documento
2. Identifica TODAS las funcionalidades mencionadas (mÃ­nimo 10 funcionalidades, incluso si el documento es corto; si no hay suficientes, extrapola basÃ¡ndote en el contexto)
3. Toma en cuenta el contexto adicional de negocio si se proporciona
4. Crea una LISTA NUMERADA de funcionalidades EXCLUSIVAMENTE para el rol: {role}.
5. Ignora cualquier funcionalidad que corresponda a otros roles diferentes a {role}.

FORMATO DE RESPUESTA:
Lista de Funcionalidades Identificadas:
1. [Nombre funcionalidad] - [DescripciÃ³n breve]
2. [Nombre funcionalidad] - [DescripciÃ³n breve]
...

Al final indica: "TOTAL FUNCIONALIDADES IDENTIFICADAS: [nÃºmero]"

NO generes historias de usuario todavÃ­a, solo la lista de funcionalidades.
"""

def create_story_generation_prompt(functionalities_list, document_text, role, business_context, start_index, batch_size=5):
    """Crea prompt para generar historias de usuario por lotes."""
    end_index = min(start_index + batch_size, len(functionalities_list))
    selected_functionalities = functionalities_list[start_index:end_index]
    func_text = "\n".join([f"{i + start_index + 1}. {func}" for i, func in enumerate(selected_functionalities)])

    context_section = ""
    if business_context and business_context.strip():
        context_section = f"""
CONTEXTO ADICIONAL DE NEGOCIO (OBLIGATORIO CONSIDERAR):
{business_context}

IMPORTANTE: Las historias de usuario deben integrar TANTO la informaciÃ³n del documento COMO las consideraciones del contexto adicional.
"""

    return f"""
Eres un analista de negocios Senior. Genera historias de usuario DETALLADAS para las siguientes funcionalidades especÃ­ficas.

FUNCIONALIDADES A DESARROLLAR (Lote {start_index + 1} a {end_index}):
{func_text}

DOCUMENTO DE REFERENCIA (para contexto adicional):
{document_text[:2000]}...
{context_section}
FORMATO OBLIGATORIO para CADA funcionalidad:

```
â•”â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
HISTORIA #{start_index + 1}: [TÃ­tulo de la funcionalidad]
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

COMO: {role}
QUIERO: [funcionalidad especÃ­fica y detallada]
PARA: [beneficio de negocio claro y medible]

CRITERIOS DE ACEPTACIÃ“N:

ğŸ”¹ Escenario Principal:
   DADO que [contexto especÃ­fico]
   CUANDO [acciÃ³n concreta del usuario]
   ENTONCES [resultado esperado detallado]

ğŸ”¹ Escenario Alternativo:
   DADO que [contexto alternativo]
   CUANDO [acciÃ³n diferente]
   ENTONCES [resultado alternativo]

ğŸ”¹ Validaciones:
   DADO que [condiciÃ³n de error]
   CUANDO [acciÃ³n que genera error]
   ENTONCES [manejo de error esperado]

REGLAS DE NEGOCIO:
â€¢ [Regla especÃ­fica 1]
â€¢ [Regla especÃ­fica 2]

PRIORIDAD: [Alta/Media/Baja]
COMPLEJIDAD: [Simple/Moderada/Compleja]

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
```

IMPORTANTE: 
- Genera AL MENOS una historia por funcionalidad.
- Si se proporciona contexto adicional, usa el contexto para ENRIQUECER las historias (por ejemplo, agregando escenarios adicionales o reglas de negocio).
- ASEGÃšRATE de que el nÃºmero total de historias no sea menor que el generado sin contexto (mÃ­nimo {len(functionalities_list)} historias).
- TODAS las historias deben generarse ÃšNICAMENTE desde la perspectiva del rol **{role}**.
- Integra el contexto adicional de negocio en las reglas de negocio y criterios de aceptaciÃ³n.
- No inventes ni incluyas otros roles diferentes a {role}.
- Numera consecutivamente desde {start_index + 1}.
"""

def create_advanced_prompt(document_text, role, story_type, business_context=None):
    """Crea el prompt avanzado basado en el tipo de historia solicitada."""
    context_section = ""
    if business_context and business_context.strip():
        context_section = f"""
CONTEXTO ADICIONAL DE NEGOCIO (CRÃTICO):
{business_context}

INTEGRACIÃ“N OBLIGATORIA: Debes incorporar este contexto en:
- Los criterios de aceptaciÃ³n
- Las reglas de negocio
- Los escenarios de validaciÃ³n
- Las consideraciones de prioridad
"""

    if story_type == 'historia de usuario' or story_type == 'funcionalidad':
        # Para documentos grandes, usar estrategia de chunks
        if len(document_text) > 5000:
            return "CHUNK_PROCESSING_NEEDED"

        # Para documentos medianos/pequeÃ±os, prompt optimizado
        prompt = f"""
Eres un analista de negocios Senior especializado en QA y anÃ¡lisis exhaustivo de requerimientos.

DOCUMENTO A ANALIZAR:
{document_text}
{context_section}
INSTRUCCIONES CRÃTICAS:

1. ANÃLISIS EXHAUSTIVO:
   - Identifica TODAS las funcionalidades del documento (mÃ­nimo 10 funcionalidades, incluso si el documento es corto; extrapola si es necesario)
   - Incluye ÃšNICAMENTE las que correspondan al rol que se proporciona en la UI {role}
   - Integra el contexto adicional de negocio en cada historia

2. GENERACIÃ“N DE HISTORIAS PARA: **{role}**

FORMATO OBLIGATORIO:

```
â•”â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
HISTORIA #{chr(123)}nÃºmero{chr(125)}: [TÃ­tulo]
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

COMO: {role}
QUIERO: [funcionalidad especÃ­fica]
PARA: [beneficio de negocio]

CRITERIOS DE ACEPTACIÃ“N:

ğŸ”¹ Escenario Principal:
   DADO que [contexto]
   CUANDO [acciÃ³n]
   ENTONCES [resultado]

ğŸ”¹ Escenario Alternativo:
   DADO que [contexto alternativo]
   CUANDO [acciÃ³n diferente]
   ENTONCES [resultado alternativo]

ğŸ”¹ Validaciones:
   DADO que [error]
   CUANDO [acciÃ³n error]
   ENTONCES [manejo error]

REGLAS DE NEGOCIO:
â€¢ [Regla 1]
â€¢ [Regla 2]

PRIORIDAD: [Alta/Media/Baja]
COMPLEJIDAD: [Simple/Moderada/Compleja]

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
```

EXPECTATIVA: Genera entre 10-50 historias segÃºn el contenido del documento, asegurando que el contexto adicional incremente el detalle y no reduzca el nÃºmero de historias.

IMPORTANTE: 
- Si el documento es extenso y sientes que podrÃ­as cortarte, termina la historia actual y agrega al final:
"CONTINÃšA EN EL SIGUIENTE LOTE - FUNCIONALIDADES PENDIENTES: [lista las que faltan]"
- SIEMPRE integra el contexto adicional proporcionado en las historias generadas.
- ASEGÃšRATE de que el nÃºmero total de historias no sea menor que el generado sin contexto (mÃ­nimo 10 historias).
"""

    elif story_type == 'caracterÃ­stica':
        prompt = f"""
Eres un analista de negocios Senior especializado en requisitos no funcionales.

DOCUMENTO A ANALIZAR:
{document_text}
{context_section}
Identifica TODOS los requisitos no funcionales (rendimiento, seguridad, usabilidad, etc.) y genera historias para el rol: {role}

FORMATO:

```
â•”â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
HISTORIA NO FUNCIONAL #{chr(123)}nÃºmero{chr(125)}: [TÃ­tulo]
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

COMO: {role}
NECESITO: [requisito no funcional]
PARA: [garantizar calidad]

CRITERIOS DE ACEPTACIÃ“N:
â€¢ [Criterio medible 1]
â€¢ [Criterio medible 2]

MÃ‰TRICAS:
â€¢ [MÃ©trica objetivo]

CATEGORÃA: [Rendimiento/Seguridad/Usabilidad/etc.]
PRIORIDAD: [Alta/Media/Baja]

â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
```

EXPECTATIVA: Genera entre 10-50 historias segÃºn el contenido del documento, asegurando que el contexto adicional incremente el detalle y no reduzca el nÃºmero de historias.

IMPORTANTE: Integra el contexto adicional de negocio en los criterios y mÃ©tricas.
"""

    else:
        # Para cualquier otro tipo, usar el formato funcional por defecto
        return create_advanced_prompt(document_text, role, 'funcionalidad', business_context)

    return prompt

def process_large_document(document_text, role, story_type, business_context=None):
    """Procesa documentos grandes dividiÃ©ndolos en chunks."""
    try:
        api_key = os.getenv("GOOGLE_API_KEY")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash-latest")

        print("ğŸ“„ Documento grande detectado. Iniciando anÃ¡lisis por fases...")
        print(f"ğŸ” Debug - business_context recibido: {business_context[:200] if business_context else 'No proporcionado'}...")
        print(f"ğŸ” Debug - role: {role}")
        print(f"ğŸ” Debug - story_type: {story_type}")

        # Fase 1: AnÃ¡lisis de funcionalidades
        print("ğŸ” Fase 1: Identificando todas las funcionalidades...")
        analysis_prompt = create_analysis_prompt(document_text, role, business_context)
        analysis_response = model.generate_content(analysis_prompt, request_options={"timeout": 90})

        # Extraer lista de funcionalidades
        functionalities = [line.strip() for line in analysis_response.text.split('\n') if re.match(r'^\d+\.', line.strip())]
        print(f"âœ… Identificadas {len(functionalities)} funcionalidades")

        # Validar nÃºmero mÃ­nimo de funcionalidades
        MIN_FUNCTIONALITIES = 10
        if len(functionalities) < MIN_FUNCTIONALITIES:
            print(f"âš ï¸ Solo se identificaron {len(functionalities)} funcionalidades, intentando generar mÃ¡s...")
            extra_prompt = analysis_prompt + f"\nINSTRUCCIÃ“N ADICIONAL: Genera al menos {MIN_FUNCTIONALITIES} funcionalidades, extrapolando si es necesario."
            extra_response = model.generate_content(extra_prompt, request_options={"timeout": 90})
            extra_functionalities = [line.strip() for line in extra_response.text.split('\n') if re.match(r'^\d+\.', line.strip())]
            functionalities.extend(extra_functionalities[:MIN_FUNCTIONALITIES - len(functionalities)])
            print(f"âœ… Total funcionalidades tras reintento: {len(functionalities)}")

        # Fase 2: Generar historias por lotes
        all_stories = []
        batch_size = max(5, len(functionalities) // 2)  # Ajustar batch_size dinÃ¡micamente
        total_batches = (len(functionalities) + batch_size - 1) // batch_size

        for batch_num in range(total_batches):
            start_idx = batch_num * batch_size
            print(f"ğŸ”¨ Generando lote {batch_num + 1}/{total_batches} (funcionalidades {start_idx + 1}-{min(start_idx + batch_size, len(functionalities))})")
            story_prompt = create_story_generation_prompt(functionalities, document_text, role, business_context, start_idx, batch_size)
            try:
                story_response = model.generate_content(story_prompt, request_options={"timeout": 120})
                all_stories.append(story_response.text)
                print(f"âœ… Lote {batch_num + 1} completado")
            except Exception as e:
                print(f"âš ï¸ Error en lote {batch_num + 1}: {e}")
                continue

        # Validar nÃºmero mÃ­nimo de historias
        MIN_STORIES = 5
        story_count = sum(story.count("HISTORIA #") for story in all_stories)
        if story_count < MIN_STORIES:
            print(f"âš ï¸ Solo se generaron {story_count} historias, intentando generar mÃ¡s...")
            extra_start_idx = len(functionalities)
            extra_prompt = create_story_generation_prompt(functionalities, document_text, role, business_context, 0, MIN_STORIES - story_count)
            extra_response = model.generate_content(extra_prompt, request_options={"timeout": 120})
            all_stories.append(extra_response.text)
            print(f"âœ… Historias adicionales generadas")

        # Combinar todas las historias
        context_summary = ""
        if business_context and business_context.strip():
            if not business_context.startswith("AIza"):
                context_summary = f"""
CONTEXTO ADICIONAL APLICADO:
{business_context[:200]}{'...' if len(business_context) > 200 else ''}
{'-' * 70}
"""
            else:
                print("âš ï¸ ADVERTENCIA: Se detectÃ³ API key en business_context, ignorando...")
                context_summary = f"""
CONTEXTO ADICIONAL APLICADO: No proporcionado
{'-' * 70}
"""
        else:
            context_summary = f"""
CONTEXTO ADICIONAL APLICADO: No proporcionado
{'-' * 70}
"""

        final_content = f"""
ANÃLISIS COMPLETO - {len(functionalities)} FUNCIONALIDADES IDENTIFICADAS
{"=" * 70}
{context_summary}
FUNCIONALIDADES IDENTIFICADAS:
{chr(10).join(functionalities)}

{"=" * 70}
HISTORIAS DE USUARIO DETALLADAS
{"=" * 70}

{chr(10).join(all_stories)}

{"=" * 70}
RESUMEN FINAL
{"=" * 70}
âœ… Total de funcionalidades procesadas: {len(functionalities)}
âœ… Total de lotes generados: {total_batches}
âœ… Contexto adicional: {'Aplicado' if business_context and not business_context.startswith("AIza") else 'No proporcionado'}
âœ… AnÃ¡lisis completado exitosamente
"""

        print("ğŸ‰ AnÃ¡lisis completo finalizado exitosamente")
        return {"status": "success", "story": final_content}

    except Exception as e:
        print(f"âŒ Error en procesamiento por chunks: {e}")
        return {"status": "error", "message": f"Error en procesamiento avanzado: {e}"}

def generate_story_from_chunk(chunk, role, story_type, business_context=None):
    """
    Genera una historia de usuario a partir de un fragmento de texto usando la API de Gemini.
    VersiÃ³n mejorada con prompts avanzados y contexto de negocio.
    """
    try:
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            return {"status": "error", "message": "API Key no configurada."}

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash-latest")

        # Crear prompt avanzado y detectar si necesita procesamiento especial
        prompt = create_advanced_prompt(chunk, role, story_type, business_context)

        # Si el documento requiere procesamiento por chunks
        if prompt == "CHUNK_PROCESSING_NEEDED":
            return process_large_document(chunk, role, story_type, business_context)

        # Generar contenido con el prompt avanzado
        response = model.generate_content(prompt, request_options={"timeout": 90})

        # Limpiar la respuesta
        story_text = response.text.strip()

        # Verificar si la respuesta se cortÃ³
        if "La generaciÃ³n completa" in story_text or "Este ejemplo ilustra" in story_text:
            print("âš ï¸ Respuesta posiblemente incompleta detectada")

        return {"status": "success", "story": story_text}

    except Exception as e:
        return {"status": "error", "message": f"Error en la generaciÃ³n: {e}"}

def create_word_document(stories):
    """Crea un documento de Word en memoria con las historias generadas."""
    doc = docx.Document()

    # TÃ­tulo principal
    title = doc.add_heading('Historias de Usuario Generadas', level=1)

    # Agregar cada historia
    for i, story in enumerate(stories, 1):
        if "HISTORIA #" in story or "â•" in story:
            doc.add_paragraph(story)
        else:
            doc.add_heading(f'Historia #{i}', level=2)
            doc.add_paragraph(story)
        doc.add_paragraph()
        doc.add_paragraph("â”€" * 50)
        doc.add_paragraph()

    return doc

def generate_story_from_text(text, role, story_type, business_context=None):
    """
    FunciÃ³n wrapper para mantener compatibilidad con la API existente
    pero usando el nuevo sistema de chunks mejorado con contexto de negocio.
    """
    chunks = split_document_into_chunks(text)
    stories = []

    for chunk in chunks:
        result = generate_story_from_chunk(chunk, role, story_type, business_context)
        if result['status'] == 'success':
            stories.append(result['story'])
        else:
            return result  # Retorna el error

    return {"status": "success", "stories": stories}

def generate_stories_with_context(document_text, role, story_type, business_context=None):
    """
    FunciÃ³n principal para generar historias de usuario con contexto de negocio.
    """
    return generate_story_from_text(document_text, role, story_type, business_context)
