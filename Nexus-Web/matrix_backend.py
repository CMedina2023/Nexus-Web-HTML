import os
import google.generativeai as genai
import docx
from pypdf import PdfReader
import csv
import json
import re
import io
import zipfile
from datetime import datetime
from difflib import SequenceMatcher
import pandas as pd


# ----------------------------
# Utilidades de lectura
# ----------------------------
def similarity(a, b):
    """Calcula similitud entre dos strings (ratio 0-1)."""
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()


def deduplicate_cases(cases):
    """Elimina casos duplicados de manera más inteligente."""
    if not cases:
        return cases

    unique_cases = []
    seen_patterns = set()

    for case in cases:
        # Crear un patrón único basado en título + tipo + categoría
        title = case['titulo_caso_prueba'].lower().strip()
        test_type = case['Tipo_de_prueba'].lower()
        category = case['Categoria'].lower() if case.get('Categoria') else ''

        # Normalizar el título (remover stopwords, puntuación)
        normalized_title = re.sub(r'\b(el|la|de|en|a|por|para|con|verificar|validar|comprobar)\b', '', title)
        normalized_title = re.sub(r'[^\w\s]', '', normalized_title).strip()

        # Crear huella digital del caso
        case_fingerprint = f"{normalized_title}_{test_type}_{category}"

        # Verificar similitud con casos existentes
        is_duplicate = False
        for seen_fingerprint in seen_patterns:
            if similarity(case_fingerprint, seen_fingerprint) > 0.85:
                is_duplicate = True
                break

        if not is_duplicate:
            seen_patterns.add(case_fingerprint)
            unique_cases.append(case)

    return unique_cases


def normalize_matrix_data(matrix_data):
    """Normaliza los datos de la matriz para consistencia."""
    normalized_data = []

    for i, case in enumerate(matrix_data, 1):
        # Crear una copia del caso para no modificar el original
        normalized_case = case.copy()

        # ASIGNAR NUEVO ID SECUENCIAL (sobreescribir cualquier ID existente)
        normalized_case['id_caso_prueba'] = f"TC{i:03d}"

        # Normalizar Pasos (siempre array)
        pasos = normalized_case.get('Pasos', [])
        if not isinstance(pasos, list):
            if isinstance(pasos, str):
                steps = []
                lines = pasos.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        # Remover numeración si existe (1., 2., etc.)
                        line = re.sub(r'^\d+[\.\)]\s*', '', line)
                        steps.append(line)
                normalized_case['Pasos'] = steps if steps else ['Paso por definir']
            else:
                normalized_case['Pasos'] = ['Paso por definir']
        else:
            # Limpiar cada paso si ya es array
            cleaned_steps = []
            for step in pasos:
                if isinstance(step, str):
                    step = step.strip()
                    step = re.sub(r'^\d+[\.\)]\s*', '', step)
                    if step:
                        cleaned_steps.append(step)
                elif step:
                    cleaned_steps.append(str(step))
            normalized_case['Pasos'] = cleaned_steps if cleaned_steps else ['Paso por definir']

        # Normalizar Resultado_esperado (siempre array)
        resultados = normalized_case.get('Resultado_esperado', [])
        if not isinstance(resultados, list):
            if isinstance(resultados, str):
                results = []
                # Separar por puntos, saltos de línea o números
                lines = re.split(r'[\.\n]|\d+[\.\)]\s*', resultados)
                for line in lines:
                    line = line.strip()
                    if line:
                        if not line.endswith('.'):
                            line += '.'
                        results.append(line)
                normalized_case['Resultado_esperado'] = results if results else ['Resultado por definir']
            else:
                normalized_case['Resultado_esperado'] = ['Resultado por definir']
        else:
            # Limpiar cada resultado si ya es array
            cleaned_results = []
            for result in resultados:
                if isinstance(result, str):
                    result = result.strip()
                    if result:
                        if not result.endswith('.'):
                            result += '.'
                        cleaned_results.append(result)
                elif result:
                    result_str = str(result).strip()
                    if result_str:
                        if not result_str.endswith('.'):
                            result_str += '.'
                        cleaned_results.append(result_str)
            normalized_case['Resultado_esperado'] = cleaned_results if cleaned_results else ['Resultado por definir']

        # Asegurar campos requeridos
        required_fields = {
            'titulo_caso_prueba': 'Título por definir',
            'Descripcion': 'Descripción por definir',
            'Precondiciones': 'Precondiciones por definir',
            'Tipo_de_prueba': 'Funcional',
            'Nivel_de_prueba': 'UAT',
            'Tipo_de_ejecucion': 'Manual',
            'Categoria': 'Flujo Principal',
            'Ambiente': 'QA',
            'Ciclo': 'Ciclo 1',
            'issuetype': 'Test Case',
            'Prioridad': 'Media',
            'historia_de_usuario': 'Historia de usuario general'
        }

        for field, default_value in required_fields.items():
            if field not in normalized_case or not normalized_case[field]:
                normalized_case[field] = default_value

        normalized_data.append(normalized_case)

    return normalized_data

def extract_text_from_file(file_path):
    """Extrae texto de archivos .docx o .pdf."""
    try:
        if file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            return "\n".join(full_text)
        elif file_path.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                text = ""
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
                return text.strip()
        else:
            raise ValueError("Formato de archivo no soportado. Usa .docx o .pdf.")
    except Exception as e:
        print(f"Error extrayendo texto del archivo: {e}")
        return ""

def split_document_into_chunks(text, max_chunk_size=4000):
    """Divide un texto largo en fragmentos más pequeños basados en la longitud."""
    if not text or len(text.strip()) == 0:
        return [""]
    if len(text) <= max_chunk_size:
        return [text]
    chunks = []
    current_chunk = ""
    paragraphs = text.split('\n')
    for paragraph in paragraphs:
        if len(paragraph) > max_chunk_size:
            sentences = re.split(r'(?<=[.!?])\s+', paragraph)
            for sentence in sentences:
                if len(current_chunk) + len(sentence) + 1 < max_chunk_size:
                    current_chunk += sentence + " "
                else:
                    if current_chunk.strip():
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence + " "
        else:
            if len(current_chunk) + len(paragraph) + 1 < max_chunk_size:
                current_chunk += paragraph + "\n"
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n"
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    if not chunks:
        chunks = [text]
    return chunks


def clean_json_response(response_text):
    """Limpia y extrae JSON de la respuesta del modelo con mejor manejo de errores."""
    if not response_text:
        return None

    # Intentar limpiar texto problemático
    cleaned_text = response_text.strip()

    # Remover markdown code blocks
    cleaned_text = re.sub(r'```json\s*', '', cleaned_text)
    cleaned_text = re.sub(r'```\s*', '', cleaned_text)

    # Buscar JSON array
    json_patterns = [
        r'\[\s*\{[\s\S]*?\}\s*\]',  # Array de objetos
        r'\{[\s\S]*?"test_cases"[\s\S]*?\}',  # Objeto con test_cases
        r'\{[\s\S]*?"matrix"[\s\S]*?\}'  # Objeto con matrix
    ]

    for pattern in json_patterns:
        match = re.search(pattern, cleaned_text, re.MULTILINE)
        if match:
            json_str = match.group(0).strip()
            try:
                data = json.loads(json_str)
                if isinstance(data, list):
                    return data
                elif isinstance(data, dict) and 'matrix' in data:
                    return data['matrix']
                elif isinstance(data, dict) and 'test_cases' in data:
                    return data['test_cases']
            except json.JSONDecodeError as e:
                print(f"Error parsing JSON: {e}")
                continue

    # Último intento: buscar cualquier array
    try:
        start = cleaned_text.find('[')
        end = cleaned_text.rfind(']')
        if start != -1 and end != -1 and end > start:
            json_str = cleaned_text[start:end + 1]
            data = json.loads(json_str)
            if isinstance(data, list):
                return data
    except json.JSONDecodeError:
        pass

    return None

def clean_text(text):
    """Limpia el texto eliminando caracteres problemáticos."""
    text = re.sub(r'[^\x20-\x7E\n]', '', text)
    return text.strip()

def extract_stories_from_text(text):
    """Extrae nombres de historias de usuario del texto del documento."""
    pattern = r'HISTORIA #\d+:[^\n]+'
    matches = re.findall(pattern, text, re.MULTILINE)
    return matches if matches else ['Historia de usuario general']

def generar_matriz_test(contexto, flujo, historia, texto_documento, tipos_prueba=['funcional', 'no_funcional']):
    try:
        api_key = "AIzaSyCAvd1ItJzSVGBL-zmHV6UkqPphW55EDlg"
        if not api_key:
            return {"status": "error",
                    "message": "API Key no configurada. Configura GEMINI_API_KEY como variable de entorno."}

        if not texto_documento or len(texto_documento.strip()) < 50:
            return {"status": "error",
                    "message": "El documento parece estar vacío o es demasiado corto. Verifica que el archivo contenga texto legible."}

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash-latest")

        # Definir prompt_base
        prompt_base = """
Eres un experto en Testing y Quality Assurance. Tu tarea es analizar requerimientos y generar casos de prueba completos.

**INSTRUCCIONES CRÍTICAS:**
1. **CONSISTENCIA**: Genera aproximadamente el mismo número de casos para historias similares
2. **COBERTURA**: Prioriza variedad sobre cantidad (no tengas un maximo ni minimo de casos de prueba, genera los necesarios para considerar una cobertura completa)
3. **EVITA DUPLICADOS**: No repitas el mismo escenario con variaciones menores
4. **ENFOQUE**: Cada caso debe cubrir un aspecto único del requerimiento
5. **FORMATO ESTRICTO**: Sigue EXACTAMENTE el formato especificado

RESPUESTA REQUERIDA: Devuelve ÚNICAMENTE un array JSON válido. Cada objeto debe tener EXACTAMENTE estas claves:

{
  "id_caso_prueba": "TC001",
  "titulo_caso_prueba": "Título descriptivo del caso",
  "Descripcion": "Descripción detallada del caso de prueba",
  "Precondiciones": "Requisitos previos para ejecutar la prueba",
  "Tipo_de_prueba": "Funcional" o "No Funcional",
  "Nivel_de_prueba": "UAT",
  "Tipo_de_ejecucion": "Manual",
  "Pasos": ["Paso 1", "Paso 2", "Paso 3"],
  "Resultado_esperado": ["Resultado esperado 1", "Resultado esperado 2"],
  "Categoria": "Categoría según el tipo de prueba",
  "Ambiente": "QA",
  "Ciclo": "Ciclo 1",
  "issuetype": "Test Case",
  "Prioridad": "Alta/Media/Baja",
  "historia_de_usuario": "Referencia a la historia de usuario"
}

**IMPORTANTE ABSOLUTO:**
- "Pasos" debe ser SIEMPRE un array de strings, SIN numeración interna
- "Resultado_esperado" debe ser SIEMPRE un array de strings completos
- NO incluir "id_caso_prueba" en tu respuesta, el sistema lo generará automáticamente
- Responde SOLO con el array JSON válido, sin texto adicional

**INSTRUCCIONES DE NUMERACIÓN:**
- NO incluir números en los títulos de los casos
- NO usar "Caso 1", "Caso 2", etc. en los títulos
- Los IDs serán generados automáticamente por el sistema

CATEGORÍAS VÁLIDAS:
- Funcional: "Flujo Principal", "Flujos Alternativos", "Casos Límite", "Casos de Error"
- No Funcional: "Rendimiento", "Seguridad", "Usabilidad", "Compatibilidad", "Confiabilidad"

IMPORTANTE: Responde SOLO con el array JSON, sin texto adicional antes o después.
        """

        # Construir prompt específico según tipos seleccionados
        incluir_funcionales = "funcional" in tipos_prueba
        incluir_no_funcionales = "no_funcional" in tipos_prueba

        if incluir_funcionales and incluir_no_funcionales:
            prompt_tipos = """
GENERA CASOS FUNCIONALES Y NO FUNCIONALES:

FUNCIONALES (no tengas un limite de casos generados, siempre y cuando el documento se preste para hacerlo):
- Flujos principales y alternativos
- Validaciones de campos y datos
- Casos límite y condiciones borde
- Manejo de errores y excepciones

NO FUNCIONALES (no tengas un limite de casos generados, siempre y cuando el documento se preste para hacerlo):
- Rendimiento y carga
- Seguridad y autorización
- Usabilidad y experiencia de usuario
- Compatibilidad entre sistemas
- Confiabilidad y disponibilidad
            """
        elif incluir_funcionales:
            prompt_tipos = """
GENERA SOLO CASOS FUNCIONALES (no tengas un limite de casos generados, siempre y cuando el documento se preste para hacerlo):
- Todos los flujos principales
- Flujos alternativos y de excepción
- Validación exhaustiva de datos
- Casos límite y condiciones extremas
- Manejo completo de errores
- Estados del sistema y transiciones
            """
        else:
            prompt_tipos = """
GENERA SOLO CASOS NO FUNCIONALES (no tengas un limite de casos generados, siempre y cuando el documento se preste para hacerlo):
- Rendimiento bajo diferentes cargas
- Seguridad y vectores de ataque
- Usabilidad en diferentes contextos
- Compatibilidad con múltiples entornos
- Confiabilidad y recuperación ante fallos
            """

        # Extraer historias del documento
        historias = extract_stories_from_text(texto_documento)
        print(f"Historias encontradas: {historias}")

        # Dividir el documento por historias en lugar de solo por tamaño
        chunks = []
        current_chunk = ""
        current_historia = historias[0] if historias else "Historia de usuario general"
        historia_chunks = {current_historia: []}

        paragraphs = texto_documento.split('\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if re.match(r'HISTORIA #\d+:', para):
                if current_chunk.strip():
                    chunks.append((current_historia, current_chunk.strip()))
                    historia_chunks[current_historia].append(current_chunk.strip())
                current_historia = para
                if current_historia not in historia_chunks:
                    historia_chunks[current_historia] = []  # Inicializar la nueva historia
                current_chunk = para + "\n"
            else:
                if len(current_chunk) + len(para) + 1 < 2500:
                    current_chunk += para + "\n"
                else:
                    chunks.append((current_historia, current_chunk.strip()))
                    historia_chunks[current_historia].append(current_chunk.strip())
                    current_chunk = para + "\n"

        if current_chunk.strip():
            chunks.append((current_historia, current_chunk.strip()))
            historia_chunks[current_historia].append(current_chunk.strip())

        print(f"Total de chunks generados: {len(chunks)}")
        print(f"Contenido de historia_chunks: {list(historia_chunks.keys())}")

        all_cases = []
        total_chunks = len(chunks)

        print(f"Procesando {total_chunks} fragmentos del documento...")

        for i, (historia_chunk, chunk) in enumerate(chunks):
            if not chunk.strip():
                print(f"Fragmento {i + 1}/{total_chunks} está vacío, omitiendo...")
                continue

            print(f"Procesando fragmento {i + 1}/{total_chunks} (Historia: {historia_chunk})")
            prompt_completo = f"{prompt_base}\n\n{prompt_tipos}\n\nCONTEXTO DEL SISTEMA: {contexto}\n\nFLUJOS A CONSIDERAR: {flujo}\n\nHISTORIA DE USUARIO: {historia}\n\nTEXTO DEL DOCUMENTO (REQUERIMIENTOS): {chunk}\n\nGenera casos de prueba basados en este requerimiento específico."

            try:
                response = model.generate_content(prompt_completo)
                if response.text.strip():
                    print(f"Respuesta del modelo para fragmento {i + 1}: {response.text[:200]}...")
                    cases_chunk = clean_json_response(response.text)
                    if cases_chunk:
                        # NORMALIZAR Y ASIGNAR IDs ÚNICOS
                        for case in cases_chunk:
                            # NORMALIZAR FORMATO DE PASOS (siempre array)
                            if not isinstance(case.get('Pasos'), list):
                                if isinstance(case.get('Pasos'), str):
                                    # Convertir string numerado a array
                                    steps = []
                                    lines = case['Pasos'].split('\n')
                                    for line in lines:
                                        line = line.strip()
                                        if line and any(char.isdigit() for char in line[:3]):
                                            # Remover numeración si existe
                                            step_text = re.sub(r'^\d+[\.\)]\s*', '', line)
                                            if step_text:
                                                steps.append(step_text)
                                        elif line:
                                            steps.append(line)
                                    case['Pasos'] = steps if steps else ['Paso por definir']
                                else:
                                    case['Pasos'] = ['Paso por definir']

                            # NORMALIZAR FORMATO DE RESULTADOS (siempre array)
                            if not isinstance(case.get('Resultado_esperado'), list):
                                if isinstance(case.get('Resultado_esperado'), str):
                                    # Convertir string a array (separar por puntos o saltos de línea)
                                    results = []
                                    lines = case['Resultado_esperado'].split('.')
                                    for line in lines:
                                        line = line.strip()
                                        if line and not line.endswith('.'):
                                            line += '.'
                                        if line:
                                            results.append(line)
                                    case['Resultado_esperado'] = results if results else ['Resultado por definir']
                                else:
                                    case['Resultado_esperado'] = ['Resultado por definir']

                            # Asegurar que la historia de usuario se asigne correctamente
                            case['historia_de_usuario'] = historia_chunk

                        all_cases.extend(cases_chunk)
                    else:
                        print(f"No se pudo procesar JSON del fragmento {i + 1}: {response.text[:500]}...")
                else:
                    print(f"Respuesta vacía del modelo para fragmento {i + 1}")
            except Exception as e:
                print(f"Error procesando fragmento {i + 1}: {str(e)}")
                continue

        # Deduplicar casos
        all_cases = deduplicate_cases(all_cases)
        all_cases = normalize_matrix_data(all_cases)
        print(
            f"Casos después de deduplicación: {len(all_cases)}")

        if not all_cases:
            return {
                "status": "error",
                "message": "No se pudieron generar casos de prueba. Verifica que el documento contenga información clara sobre requerimientos o funcionalidades."
            }

        funcional_count = sum(1 for case in all_cases if case.get('Tipo_de_prueba', '').lower() == 'funcional')
        no_funcional_count = len(all_cases) - funcional_count

        return {
            "status": "success",
            "matrix": all_cases,
            "total_cases": len(all_cases),
            "funcional_cases": funcional_count,
            "no_funcional_cases": no_funcional_count
        }
    except Exception as e:
        error_message = str(e).lower()
        print(f"Error general: {str(e)}")
        print(f"Tipo de excepción: {type(e).__name__}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        if "blocked" in error_message or "safety" in error_message:
            return {
                "status": "error",
                "message": "La solicitud fue bloqueada por filtros de seguridad. Intenta con un documento diferente."
            }
        else:
            return {
                "status": "error",
                "message": f"Error en la lógica de procesamiento: {str(e)}"
            }

def save_to_xlsx_buffer(data):
    """Guarda los datos de la matriz en un buffer de memoria como XLSX."""
    if not data:
        return b""

    # Campos en el orden deseado (igual que CSV)
    fieldnames = [
        "id_caso_prueba",
        "titulo_caso_prueba",
        "Descripcion",
        "Precondiciones",
        "Tipo_de_prueba",
        "Nivel_de_prueba",
        "Tipo_de_ejecucion",
        "Pasos",
        "Resultado_esperado",
        "Categoria",
        "Ambiente",
        "Ciclo",
        "issuetype",
        "Prioridad",
        "historia_de_usuario"
    ]

    # Crear DataFrame con pandas
    df = pd.DataFrame(data)[fieldnames]  # Selecciona solo las columnas en orden

    # Buffer para XLSX
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Matriz de Pruebas', index=False)
        # Opcional: Auto-ajustar columnas para mejor legibilidad
        worksheet = writer.sheets['Matriz de Pruebas']
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)  # Límite para no hacer columnas eternas
            worksheet.column_dimensions[column_letter].width = adjusted_width

    output.seek(0)
    return output.getvalue()

def save_to_csv_buffer(data):
    """Guarda los datos de la matriz en un buffer de memoria como CSV."""
    if not data:
        return b""

    # Campos en el orden deseado
    fieldnames = [
        "id_caso_prueba",
        "titulo_caso_prueba",
        "Descripcion",
        "Precondiciones",
        "Tipo_de_prueba",
        "Nivel_de_prueba",
        "Tipo_de_ejecucion",
        "Pasos",
        "Resultado_esperado",
        "Categoria",
        "Ambiente",
        "Ciclo",
        "issuetype",
        "Prioridad",
        "historia_de_usuario"
    ]

    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()

    for row in data:
        # Crear una copia del row para no modificar el original
        csv_row = {}

        for field in fieldnames:
            value = row.get(field, '')

            # Convertir listas a string separado por " | "
            if field in ['Pasos', 'Resultado_esperado'] and isinstance(value, list):
                csv_row[field] = " | ".join(str(item) for item in value if item)
            elif isinstance(value, list):
                csv_row[field] = ", ".join(str(item) for item in value if item)
            else:
                csv_row[field] = str(value) if value else ''

        writer.writerow(csv_row)

    return output.getvalue().encode('utf-8')


def save_to_json_buffer(data):
    """Guarda los datos de la matriz en un buffer de memoria como JSON."""
    if not data:
        return b"[]"

    output = io.StringIO()
    json.dump(data, output, indent=4, ensure_ascii=False)
    return output.getvalue().encode('utf-8')


def create_zip_with_matrix(data, output_filename):
    """
    Crea un archivo ZIP con la matriz en formato CSV, JSON y XLSX.
    """
    if not data:
        return None

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Agregar archivo CSV (igual que antes)
        csv_data = save_to_csv_buffer(data)
        zip_file.writestr(f"{output_filename}.csv", csv_data)

        # Agregar archivo JSON (igual que antes)
        json_data = save_to_json_buffer(data)
        zip_file.writestr(f"{output_filename}.json", json_data)

        # NUEVO: Agregar archivo XLSX
        xlsx_data = save_to_xlsx_buffer(data)
        zip_file.writestr(f"{output_filename}.xlsx", xlsx_data)

        # Agregar archivo README actualizado con mención al XLSX
        funcional_count = sum(1 for case in data if case.get('Tipo_de_prueba', '').lower() == 'funcional')
        no_funcional_count = len(data) - funcional_count
        readme_content = f"""MATRIZ DE PRUEBAS GENERADA
============================

Archivo generado automáticamente por Matrix Generator
Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

CONTENIDO DEL ZIP:
- {output_filename}.csv: Matriz de pruebas en formato CSV (texto plano, ideal para importaciones simples)
- {output_filename}.json: Matriz de pruebas en formato JSON (para APIs e integraciones)
- {output_filename}.xlsx: Matriz de pruebas en formato Excel (con hoja 'Matriz de Pruebas', columnas auto-ajustadas)
- README.txt: Este archivo

ESTADÍSTICAS:
- Total de casos de prueba: {len(data)}
- Casos funcionales: {funcional_count}
- Casos no funcionales: {no_funcional_count}

ESTRUCTURA DE CAMPOS:
- id_caso_prueba: Identificador único del caso
- titulo_caso_prueba: Título descriptivo
- Descripcion: Descripción detallada del caso
- Precondiciones: Requisitos previos
- Tipo_de_prueba: Funcional o No Funcional
- Nivel_de_prueba: Nivel de testing (UAT)
- Tipo_de_ejecucion: Manual o Automático
- Pasos: Pasos a seguir (separados por " | " en CSV/XLSX)
- Resultado_esperado: Resultados esperados (separados por " | " en CSV/XLSX)
- Categoria: Categoría específica del tipo de prueba
- Ambiente: Ambiente de pruebas (QA)
- Ciclo: Ciclo de testing
- issuetype: Tipo de issue (Test Case)
- Prioridad: Alta, Media o Baja
- historia_de_usuario: Referencia a la historia

INSTRUCCIONES DE USO:
1. Abre el XLSX en Excel para ver/filtrar/editar fácilmente.
2. Importa el CSV en herramientas como Jira, TestRail o Google Sheets.
3. Usa el JSON para scripts automatizados.
4. Revisa y ajusta los casos según tus necesidades específicas.
"""
        zip_file.writestr("README.txt", readme_content.encode('utf-8'))

    zip_buffer.seek(0)
    return zip_buffer.getvalue()


def process_matrix_request(file_path, contexto="", flujo="", historia="", tipos_prueba=['funcional', 'no_funcional'],
                           output_filename="matriz_pruebas"):
    """
    Función principal que procesa una solicitud completa de generación de matriz.

    Args:
        file_path (str): Ruta al archivo de requerimientos
        contexto (str): Contexto del sistema
        flujo (str): Flujo específico a probar
        historia (str): Historia de usuario
        tipos_prueba (list): Tipos de pruebas a generar
        output_filename (str): Nombre base para archivos de salida

    Returns:
        dict: Resultado de la operación
    """
    try:
        # Extraer texto del documento
        print(f"Extrayendo texto del archivo: {file_path}")
        texto_documento = extract_text_from_file(file_path)

        if not texto_documento or len(texto_documento.strip()) < 50:
            return {
                "status": "error",
                "message": "No se pudo extraer texto del documento o el contenido es insuficiente."
            }

        print(f"Texto extraído: {len(texto_documento)} caracteres")

        # Generar matriz de pruebas
        print("Generando matriz de pruebas...")
        result = generar_matriz_test(contexto, flujo, historia, texto_documento, tipos_prueba)

        if result["status"] != "success":
            return result

        # Crear archivo ZIP
        print("Creando archivo ZIP...")
        zip_data = create_zip_with_matrix(result["matrix"], output_filename)

        if not zip_data:
            return {
                "status": "error",
                "message": "Error al crear el archivo ZIP."
            }

        return {
            "status": "success",
            "message": "Matriz generada exitosamente",
            "zip_data": zip_data,
            "stats": {
                "total_cases": result.get("total_cases", 0),
                "funcional_cases": result.get("funcional_cases", 0),
                "no_funcional_cases": result.get("no_funcional_cases", 0)
            }
        }

    except Exception as e:
        return {
            "status": "error",
            "message": f"Error procesando la solicitud: {str(e)}"
        }

def extract_stories_from_text(text):
    """Extrae nombres de historias de usuario del texto del documento."""
    pattern = r'HISTORIA #\d+:[^\n]+'
    matches = re.findall(pattern, text, re.MULTILINE)
    return matches if matches else ['Historia de usuario general']


def test_matrix_generation():
    """
    Función de prueba para verificar la generación de matrices.
    """
    # Texto de prueba
    texto_prueba = """
    Requerimiento: Sistema de Login de Usuario

    El sistema debe permitir a los usuarios autenticarse usando email y contraseña.

    Funcionalidades:
    1. Campo de email con validación de formato
    2. Campo de contraseña con mínimo 8 caracteres
    3. Botón de "Iniciar Sesión"
    4. Opción "Recordar usuario"
    5. Link "Olvidé mi contraseña"
    6. Mensaje de error para credenciales inválidas
    7. Redirección al dashboard después del login exitoso

    Reglas de negocio:
    - Después de 3 intentos fallidos, bloquear la cuenta por 15 minutos
    - La sesión debe expirar después de 2 horas de inactividad
    - Debe registrar todos los intentos de login en el log de auditoría
    """

    print("Ejecutando prueba de generación de matriz...")

    result = generar_matriz_test(
        contexto="Sistema web de gestión de usuarios",
        flujo="Login de usuario con email y contraseña",
        historia="Como usuario quiero poder iniciar sesión de forma segura",
        texto_documento=texto_prueba,
        tipos_prueba=['funcional', 'no_funcional']
    )

    print(f"Resultado: {result['status']}")
    if result['status'] == 'success':
        print(f"Casos generados: {len(result['matrix'])}")
        for i, case in enumerate(result['matrix'][:3]):  # Mostrar solo los primeros 3
            print(f"\nCaso {i + 1}:")
            print(f"  ID: {case.get('id_caso_prueba', 'N/A')}")
            print(f"  Título: {case.get('titulo_caso_prueba', 'N/A')}")
            print(f"  Tipo: {case.get('Tipo_de_prueba', 'N/A')}")
    else:
        print(f"Error: {result['message']}")

if __name__ == "__main__":
    # Ejecutar prueba si se ejecuta directamente
    test_matrix_generation()
